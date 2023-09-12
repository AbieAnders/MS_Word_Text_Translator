import os
#docx library
import docx
from docx.document import Document
from docx.text.paragraph import Paragraph
from docx.enum.style import WD_STYLE_TYPE
from docx.table import _Cell, _Row, Table
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
#translatepy library
from translatepy import Translator

#Path of the desired new working directory
os.chdir('Path of the desired new working directory')
#New document
doc = docx.Document()
doc.save('Fully Translated Document 2.docx')
#Document of choice
doc1 = docx.Document('Document of choice.docx')
print("The documents were accessed successfully.")

translator_object = Translator()
def Translate_Text(file_text, language_code):
    try:
        translation = translator_object.translate(file_text,str(language_code))
        if translation is None:
            return None
        return str(translation)
    except IndexError:
        print("Index error")
        return None
    except Exception as e:
        print("An error has occurred during translation:", e)
        return None

def Iter_Block_Items(parent):
    """
    Generate a reference to each paragraph and table child within *parent*,
    in document order. Each returned value is an instance of either Table or
    Paragraph. *parent* would most commonly be a reference to a main
    Document object, but also works for a _Cell object, which itself can
    contain paragraphs and tables.
    """
    if isinstance(parent, Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    elif isinstance(parent, _Row):
        parent_elm = parent._tr
    else:
        raise ValueError("something's not right")
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent_elm)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent_elm)

#The Font style is created here since direct access of the font style is not supported by the library
def Create_Font(document, style_name):
    doc_styles = document.styles
    doc_charstyle = doc_styles.add_style('Ordinary ' + str(style_name), WD_STYLE_TYPE.CHARACTER)
    obj_font = doc_charstyle.font
    obj_font.name = str(style_name)
    return None

##Execution loop

#Style and Font info are gathered from all the paragraphs of the document(Gathering Style and Font info from Tables is not possible in most circumstances)
font_list = []
style_list = []
para_counter = 1
for para in doc1.paragraphs:
    print("block count:", para_counter)
    #print(para.text)
    if para.text.strip() == '':
        print(para_counter)
        para_counter += 1
        continue
    font_list.append(para.style.font.name)
    style_list.append(para.style.name)
    para_counter += 1

#This section removes all repeating Font types and the 'None' type font from font_list and stores it in a new variable.
set_font_list = set(font_list)
try:
    set_font_list.remove(None)
except:
    print("The document does not have any 'None' type fonts.")
set_font_list = list(set_font_list)
#A set is an unordered collection of unique elements and so type casting a list that has repeating elements into a set removes the repeating elements.
#The set can then be type casted into a list once again.
set_style_list = list(set(style_list))  

#print("Number of Fonts:", len(font_list))
print(font_list)
print(set_font_list)
#print("Number of Styles:", len(style_list))
print(style_list)
print(set_style_list)

for font in set_font_list:
    Create_Font(doc, font)
print("Execution of the Main block is ready to being.")

#Main execution
font_counter = 0
style_counter = 0
for count, block in enumerate(Iter_Block_Items(doc1)):
    count += 1
    print("block count:", count)
    if isinstance(block, Paragraph):
        #print(block.text)
        if block.text.strip() == '':
            print(count)
            continue
        if(font_list[font_counter] in set_font_list):
            #set_font_list.index(font_list[font_counter])
            translated_text = Translate_Text(block.text, 'ta')
            para = doc.add_paragraph()
            para_runner = para.add_run(translated_text, style = 'Ordinary ' + font_list[font_counter]).bold = True
            doc.save('Fully Translated Document 2.docx')
            font_counter += 1    #Used since the continue statement causes the font_counter incrementation that is found below to not execute.
            #style_counter += 1   #Unncessary since it is not being used.
            continue
        block = Translate_Text(block.text, 'ta')
        doc.add_paragraph(block)
        doc.save('Fully Translated Document 2.docx')
        font_counter += 1
        style_counter += 1
    elif isinstance(block, Table):
        n = len(block.rows)
        m = int(len(block._cells)/n)
        print(n, m)
        table = doc.add_table(n, m)
        for a,row in enumerate(block.rows):
            for b,cell in enumerate(row.cells):
                if cell.text.strip() == '':
                    print("Empty cell")
                    continue
                #print(cell.text)
                translation = translator_object.translate(cell.text,'ta')
                table.cell(a, b).text = str(translation)  #str type casting is necessary since the text function automatically assumes the TranslationResult type
        doc.save('Fully Translated Document 2.docx')

#table.style = 'Colorful List'
doc.save('Fully Translated Document 2.docx')
